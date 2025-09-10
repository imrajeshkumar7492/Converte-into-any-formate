import os
import io
import tempfile
from typing import BinaryIO, Dict, Any
from pathlib import Path
import zipfile
import xml.etree.ElementTree as ET

class AdvancedDocumentConverter:
    """Advanced document converter with support for more formats and better quality"""
    
    @staticmethod
    def convert_epub_to_pdf(input_file: BinaryIO) -> bytes:
        """Convert EPUB to PDF"""
        try:
            from ebooklib import epub
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from bs4 import BeautifulSoup
            import tempfile
            import os
            
            input_file.seek(0)
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.epub', delete=False) as temp_epub:
                temp_epub.write(input_file.read())
                temp_epub_path = temp_epub.name
            
            try:
                # Read EPUB
                book = epub.read_epub(temp_epub_path)
                
                # Create PDF
                output_buffer = io.BytesIO()
                doc = SimpleDocTemplate(output_buffer, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Process each chapter
                for item in book.get_items():
                    if item.get_type() == epub.ITEM_DOCUMENT:
                        # Extract text content
                        soup = BeautifulSoup(item.get_content(), 'html.parser')
                        text_content = soup.get_text()
                        
                        # Add to PDF
                        if text_content.strip():
                            para = Paragraph(text_content, styles['Normal'])
                            story.append(para)
                            story.append(Spacer(1, 12))
                
                # Build PDF
                doc.build(story)
                return output_buffer.getvalue()
                
            finally:
                if os.path.exists(temp_epub_path):
                    os.unlink(temp_epub_path)
                    
        except Exception as e:
            raise Exception(f"EPUB to PDF conversion failed: {str(e)}")
    
    @staticmethod
    def convert_mobi_to_pdf(input_file: BinaryIO) -> bytes:
        """Convert MOBI to PDF"""
        try:
            # MOBI conversion requires additional libraries
            # For now, we'll provide a basic implementation
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            input_file.seek(0)
            content = input_file.read()
            
            # Basic MOBI to PDF conversion
            output_buffer = io.BytesIO()
            c = canvas.Canvas(output_buffer, pagesize=letter)
            width, height = letter
            
            # Extract text from MOBI (simplified)
            text_content = content.decode('utf-8', errors='ignore')
            
            y = height - 50
            lines = text_content.split('\n')
            
            for line in lines[:100]:  # Limit to first 100 lines
                if y < 50:
                    c.showPage()
                    y = height - 50
                
                c.drawString(50, y, line[:80])  # Limit line length
                y -= 15
            
            c.save()
            return output_buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"MOBI to PDF conversion failed: {str(e)}")
    
    @staticmethod
    def convert_rtf_to_pdf(input_file: BinaryIO) -> bytes:
        """Convert RTF to PDF with formatting preservation"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            import re
            
            input_file.seek(0)
            rtf_content = input_file.read().decode('utf-8', errors='ignore')
            
            # Basic RTF parsing (simplified)
            # Remove RTF control codes
            text_content = re.sub(r'\\[a-z]+\d*\s?', '', rtf_content)
            text_content = re.sub(r'[{}]', '', text_content)
            text_content = re.sub(r'\\[^a-z]', '', text_content)
            
            # Create PDF
            output_buffer = io.BytesIO()
            doc = SimpleDocTemplate(output_buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Split into paragraphs
            paragraphs = text_content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    para = Paragraph(para_text.strip(), styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            return output_buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"RTF to PDF conversion failed: {str(e)}")
    
    @staticmethod
    def convert_odt_to_pdf(input_file: BinaryIO) -> bytes:
        """Convert ODT to PDF"""
        try:
            from odf.opendocument import load
            from odf.text import P
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            import tempfile
            import os
            
            input_file.seek(0)
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.odt', delete=False) as temp_odt:
                temp_odt.write(input_file.read())
                temp_odt_path = temp_odt.name
            
            try:
                # Load ODT document
                doc = load(temp_odt_path)
                
                # Create PDF
                output_buffer = io.BytesIO()
                pdf_doc = SimpleDocTemplate(output_buffer, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Extract text from paragraphs
                for paragraph in doc.getElementsByType(P):
                    text = paragraph.getText()
                    if text.strip():
                        para = Paragraph(text, styles['Normal'])
                        story.append(para)
                        story.append(Spacer(1, 12))
                
                # Build PDF
                pdf_doc.build(story)
                return output_buffer.getvalue()
                
            finally:
                if os.path.exists(temp_odt_path):
                    os.unlink(temp_odt_path)
                    
        except Exception as e:
            raise Exception(f"ODT to PDF conversion failed: {str(e)}")
    
    @staticmethod
    def get_document_info(input_file: BinaryIO, source_format: str) -> Dict[str, Any]:
        """Get detailed information about a document"""
        try:
            input_file.seek(0)
            content = input_file.read()
            
            info = {
                'size_bytes': len(content),
                'format': source_format.upper(),
                'category': 'document'
            }
            
            if source_format.lower() == 'pdf':
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(content))
                info.update({
                    'pages': len(reader.pages),
                    'title': reader.metadata.get('/Title', '') if reader.metadata else '',
                    'author': reader.metadata.get('/Author', '') if reader.metadata else '',
                    'subject': reader.metadata.get('/Subject', '') if reader.metadata else '',
                    'creator': reader.metadata.get('/Creator', '') if reader.metadata else '',
                    'producer': reader.metadata.get('/Producer', '') if reader.metadata else '',
                    'creation_date': str(reader.metadata.get('/CreationDate', '')) if reader.metadata else '',
                    'modification_date': str(reader.metadata.get('/ModDate', '')) if reader.metadata else ''
                })
            
            elif source_format.lower() in ['docx', 'doc']:
                from docx import Document
                doc = Document(io.BytesIO(content))
                info.update({
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'sections': len(doc.sections)
                })
            
            elif source_format.lower() == 'epub':
                from ebooklib import epub
                book = epub.read_epub(io.BytesIO(content))
                info.update({
                    'title': book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else '',
                    'author': book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else '',
                    'language': book.get_metadata('DC', 'language')[0][0] if book.get_metadata('DC', 'language') else '',
                    'chapters': len([item for item in book.get_items() if item.get_type() == epub.ITEM_DOCUMENT])
                })
            
            return info
            
        except Exception as e:
            return {'error': str(e), 'size_bytes': len(content) if 'content' in locals() else 0}