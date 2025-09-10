import os
import io
import tempfile
from typing import BinaryIO
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Inches
import openpyxl
from openpyxl.utils import get_column_letter
import csv
from pptx import Presentation
from PIL import Image
import fitz  # PyMuPDF for PDF to image conversion

class DocumentConverter:
    @staticmethod
    def convert_pdf_to_text(input_file: BinaryIO) -> bytes:
        """Convert PDF to plain text"""
        try:
            input_file.seek(0)
            reader = PdfReader(input_file)
            text_content = ""
            
            for page in reader.pages:
                text_content += page.extract_text() + "\n\n"
            
            return text_content.encode('utf-8')
        except Exception as e:
            raise Exception(f"PDF to text conversion failed: {str(e)}")
    
    @staticmethod
    def convert_pdf_to_docx(input_file: BinaryIO) -> bytes:
        """Convert PDF to DOCX"""
        try:
            input_file.seek(0)
            reader = PdfReader(input_file)
            
            # Create new document
            doc = Document()
            
            for page_num, page in enumerate(reader.pages):
                if page_num > 0:
                    doc.add_page_break()
                
                text = page.extract_text()
                doc.add_paragraph(text)
            
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            
            return output_buffer.getvalue()
        except Exception as e:
            raise Exception(f"PDF to DOCX conversion failed: {str(e)}")
    
    @staticmethod
    def convert_docx_to_pdf(input_file: BinaryIO) -> bytes:
        """Convert DOCX to PDF with advanced formatting preservation"""
        try:
            # Try using python-docx2pdf first (better formatting preservation)
            try:
                from docx2pdf import convert
                import tempfile
                import os
                
                input_file.seek(0)
                
                # Create temporary files
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                    temp_docx.write(input_file.read())
                    temp_docx_path = temp_docx.name
                
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                    temp_pdf_path = temp_pdf.name
                
                try:
                    # Convert using docx2pdf (preserves formatting better)
                    convert(temp_docx_path, temp_pdf_path)
                    
                    # Read the converted PDF
                    with open(temp_pdf_path, 'rb') as pdf_file:
                        pdf_content = pdf_file.read()
                    
                    return pdf_content
                    
                finally:
                    # Clean up temporary files
                    for path in [temp_docx_path, temp_pdf_path]:
                        if os.path.exists(path):
                            os.unlink(path)
                            
            except ImportError:
                # Fallback to advanced reportlab conversion
                return DocumentConverter._convert_docx_to_pdf_advanced(input_file)
                
        except Exception as e:
            raise Exception(f"DOCX to PDF conversion failed: {str(e)}")
    
    @staticmethod
    def _convert_docx_to_pdf_advanced(input_file: BinaryIO) -> bytes:
        """Advanced DOCX to PDF conversion with better formatting preservation"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
            from docx.shared import Inches
            import tempfile
            import os
            
            input_file.seek(0)
            doc = Document(input_file)
            
            output_buffer = io.BytesIO()
            doc_pdf = SimpleDocTemplate(output_buffer, pagesize=A4, 
                                      rightMargin=72, leftMargin=72, 
                                      topMargin=72, bottomMargin=18)
            
            # Define styles
            styles = getSampleStyleSheet()
            
            # Create custom styles based on document formatting
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=12,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=8,
                alignment=TA_LEFT
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                alignment=TA_JUSTIFY
            )
            
            # Build content
            story = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    story.append(Spacer(1, 6))
                    continue
                
                # Determine style based on paragraph properties
                if paragraph.style.name.startswith('Title'):
                    story.append(Paragraph(text, title_style))
                elif paragraph.style.name.startswith('Heading'):
                    story.append(Paragraph(text, heading_style))
                else:
                    # Check for bold/italic formatting
                    if paragraph.runs:
                        formatted_text = ""
                        for run in paragraph.runs:
                            if run.bold and run.italic:
                                formatted_text += f"<b><i>{run.text}</i></b>"
                            elif run.bold:
                                formatted_text += f"<b>{run.text}</b>"
                            elif run.italic:
                                formatted_text += f"<i>{run.text}</i>"
                            else:
                                formatted_text += run.text
                        story.append(Paragraph(formatted_text, normal_style))
                    else:
                        story.append(Paragraph(text, normal_style))
                
                story.append(Spacer(1, 6))
            
            # Handle tables
            for table in doc.tables:
                from reportlab.platypus import Table, TableStyle
                from reportlab.lib import colors
                
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    table_obj = Table(table_data)
                    table_obj.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table_obj)
                    story.append(Spacer(1, 12))
            
            # Build PDF
            doc_pdf.build(story)
            return output_buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"Advanced DOCX to PDF conversion failed: {str(e)}")
    
    @staticmethod
    def convert_docx_to_txt(input_file: BinaryIO) -> bytes:
        """Convert DOCX to plain text"""
        try:
            input_file.seek(0)
            doc = Document(input_file)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return text_content.encode('utf-8')
        except Exception as e:
            raise Exception(f"DOCX to text conversion failed: {str(e)}")
    
    @staticmethod
    def convert_txt_to_pdf(input_file: BinaryIO) -> bytes:
        """Convert plain text to PDF"""
        try:
            input_file.seek(0)
            text_content = input_file.read().decode('utf-8')
            
            output_buffer = io.BytesIO()
            c = canvas.Canvas(output_buffer, pagesize=letter)
            width, height = letter
            
            lines = text_content.split('\n')
            y = height - 50
            
            for line in lines:
                if y < 50:  # Start new page
                    c.showPage()
                    y = height - 50
                
                # Handle long lines
                if len(line) * 6 > width - 100:
                    words = line.split()
                    current_line = ""
                    for word in words:
                        test_line = current_line + word + " "
                        if len(test_line) * 6 > width - 100:
                            if current_line:
                                c.drawString(50, y, current_line.strip())
                                y -= 15
                                current_line = word + " "
                            else:
                                c.drawString(50, y, word)
                                y -= 15
                        else:
                            current_line = test_line
                    
                    if current_line.strip():
                        c.drawString(50, y, current_line.strip())
                        y -= 15
                else:
                    c.drawString(50, y, line)
                    y -= 15
            
            c.save()
            return output_buffer.getvalue()
        except Exception as e:
            raise Exception(f"Text to PDF conversion failed: {str(e)}")
    
    @staticmethod
    def convert_excel_to_csv(input_file: BinaryIO) -> bytes:
        """Convert Excel to CSV"""
        try:
            input_file.seek(0)
            workbook = openpyxl.load_workbook(input_file)
            worksheet = workbook.active
            
            output = io.StringIO()
            writer = csv.writer(output)
            
            for row in worksheet.iter_rows(values_only=True):
                # Convert None values to empty strings
                clean_row = [str(cell) if cell is not None else '' for cell in row]
                writer.writerow(clean_row)
            
            return output.getvalue().encode('utf-8')
        except Exception as e:
            raise Exception(f"Excel to CSV conversion failed: {str(e)}")
    
    @staticmethod
    def convert_csv_to_excel(input_file: BinaryIO) -> bytes:
        """Convert CSV to Excel"""
        try:
            input_file.seek(0)
            csv_content = input_file.read().decode('utf-8')
            
            workbook = openpyxl.Workbook()
            worksheet = workbook.active
            
            # Parse CSV
            csv_reader = csv.reader(io.StringIO(csv_content))
            for row_num, row in enumerate(csv_reader, 1):
                for col_num, value in enumerate(row, 1):
                    worksheet.cell(row=row_num, column=col_num, value=value)
            
            output_buffer = io.BytesIO()
            workbook.save(output_buffer)
            
            return output_buffer.getvalue()
        except Exception as e:
            raise Exception(f"CSV to Excel conversion failed: {str(e)}")
    
    @staticmethod
    def convert_pdf_to_image(input_file: BinaryIO, target_format: str = 'png') -> bytes:
        """Convert PDF to image (JPG/PNG) - supports multi-page PDFs by creating a combined image"""
        try:
            # Try using PyMuPDF (fitz) first - better for PDF to image conversion
            try:
                import fitz  # PyMuPDF
                from PIL import Image
                
                input_file.seek(0)
                pdf_data = input_file.read()
                
                # Open PDF with PyMuPDF
                pdf_document = fitz.open(stream=pdf_data, filetype="pdf")
                
                # Get all pages
                num_pages = len(pdf_document)
                if num_pages == 0:
                    raise Exception("PDF has no pages")
                
                # For single page, convert directly
                if num_pages == 1:
                    page = pdf_document[0]
                    mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_data))
                else:
                    # For multi-page PDFs, create a combined image
                    images = []
                    max_width = 0
                    total_height = 0
                    
                    for page_num in range(num_pages):
                        page = pdf_document[page_num]
                        mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")
                        page_image = Image.open(io.BytesIO(img_data))
                        images.append(page_image)
                        
                        max_width = max(max_width, page_image.width)
                        total_height += page_image.height
                    
                    # Create combined image
                    combined_image = Image.new('RGB', (max_width, total_height), 'white')
                    y_offset = 0
                    
                    for page_image in images:
                        # Center the page image if it's narrower than max_width
                        x_offset = (max_width - page_image.width) // 2
                        combined_image.paste(page_image, (x_offset, y_offset))
                        y_offset += page_image.height
                    
                    image = combined_image
                
                # Convert to target format
                output_buffer = io.BytesIO()
                
                if target_format.lower() in ['jpg', 'jpeg']:
                    # Convert to RGB for JPEG (remove alpha channel)
                    if image.mode in ['RGBA', 'LA']:
                        background = Image.new('RGB', image.size, (255, 255, 255))
                        if image.mode == 'RGBA':
                            background.paste(image, mask=image.split()[-1])
                        else:
                            background.paste(image)
                        image = background
                    elif image.mode != 'RGB':
                        image = image.convert('RGB')
                    
                    image.save(output_buffer, format='JPEG', quality=95, optimize=True)
                else:  # PNG
                    if image.mode != 'RGBA':
                        image = image.convert('RGBA')
                    image.save(output_buffer, format='PNG', optimize=True)
                
                pdf_document.close()
                return output_buffer.getvalue()
                
            except (ImportError, Exception) as e:
                # Fallback: Use reportlab to create a simple image representation
                # This is a basic fallback - not ideal but better than nothing
                from PIL import Image, ImageDraw, ImageFont
                
                input_file.seek(0)
                reader = PdfReader(input_file)
                
                # Extract text from first page
                first_page = reader.pages[0]
                text_content = first_page.extract_text()
                
                # Create a simple image with the text
                # Create a white image
                img_width, img_height = 800, 1000
                image = Image.new('RGB', (img_width, img_height), 'white')
                draw = ImageDraw.Draw(image)
                
                # Try to use a default font
                try:
                    font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12)
                except:
                    font = ImageFont.load_default()
                
                # Draw text on image (simple word wrapping)
                y_position = 50
                words = text_content.split()
                line = ""
                
                for word in words:
                    test_line = line + word + " "
                    bbox = draw.textbbox((0, 0), test_line, font=font)
                    if bbox[2] > img_width - 100:  # Line too long
                        if line:
                            draw.text((50, y_position), line.strip(), fill='black', font=font)
                            y_position += 20
                            line = word + " "
                        else:
                            draw.text((50, y_position), word, fill='black', font=font)
                            y_position += 20
                    else:
                        line = test_line
                    
                    if y_position > img_height - 100:  # Image full
                        break
                
                if line.strip():
                    draw.text((50, y_position), line.strip(), fill='black', font=font)
                
                # Save as target format
                output_buffer = io.BytesIO()
                
                if target_format.lower() in ['jpg', 'jpeg']:
                    image.save(output_buffer, format='JPEG', quality=95)
                else:  # PNG
                    image.save(output_buffer, format='PNG')
                
                return output_buffer.getvalue()
                
        except Exception as e:
            raise Exception(f"PDF to image conversion failed: {str(e)}")
    
    @staticmethod
    def convert_pdf_to_jpg(input_file: BinaryIO) -> bytes:
        """Convert PDF to JPG"""
        return DocumentConverter.convert_pdf_to_image(input_file, 'jpg')
    
    @staticmethod
    def convert_pdf_to_png(input_file: BinaryIO) -> bytes:
        """Convert PDF to PNG"""
        return DocumentConverter.convert_pdf_to_image(input_file, 'png')
    
    @staticmethod
    def convert_pdf_to_images_zip(input_file: BinaryIO, target_format: str = 'png') -> bytes:
        """Convert PDF to ZIP file containing individual page images"""
        try:
            import fitz  # PyMuPDF
            from PIL import Image
            import zipfile
            
            input_file.seek(0)
            pdf_data = input_file.read()
            
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(stream=pdf_data, filetype="pdf")
            num_pages = len(pdf_document)
            
            if num_pages == 0:
                raise Exception("PDF has no pages")
            
            # Create ZIP file in memory
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for page_num in range(num_pages):
                    page = pdf_document[page_num]
                    mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    page_image = Image.open(io.BytesIO(img_data))
                    
                    # Convert to target format
                    img_buffer = io.BytesIO()
                    
                    if target_format.lower() in ['jpg', 'jpeg']:
                        if page_image.mode in ['RGBA', 'LA']:
                            background = Image.new('RGB', page_image.size, (255, 255, 255))
                            if page_image.mode == 'RGBA':
                                background.paste(page_image, mask=page_image.split()[-1])
                            else:
                                background.paste(page_image)
                            page_image = background
                        elif page_image.mode != 'RGB':
                            page_image = page_image.convert('RGB')
                        
                        page_image.save(img_buffer, format='JPEG', quality=95, optimize=True)
                    else:  # PNG
                        if page_image.mode != 'RGBA':
                            page_image = page_image.convert('RGBA')
                        page_image.save(img_buffer, format='PNG', optimize=True)
                    
                    # Add to ZIP
                    filename = f"page_{page_num + 1}.{target_format.lower()}"
                    zip_file.writestr(filename, img_buffer.getvalue())
            
            pdf_document.close()
            return zip_buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"PDF to images ZIP conversion failed: {str(e)}")